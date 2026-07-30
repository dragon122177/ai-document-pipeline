from __future__ import annotations

import csv
import io
import json
from dataclasses import dataclass
from pathlib import Path

from docx import Document as WordDocument
from pypdf import PdfReader


class ExtractionError(ValueError):
    pass


@dataclass(slots=True)
class ExtractedContent:
    text: str
    page_count: int
    language_hint: str | None = None


def extract_file(path: Path) -> ExtractedContent:
    suffix = path.suffix.lower()
    if suffix in {".txt", ".md"}:
        return ExtractedContent(_decode_text(path.read_bytes()), 1)
    if suffix == ".csv":
        return _extract_csv(path)
    if suffix == ".json":
        return _extract_json(path)
    if suffix == ".pdf":
        return _extract_pdf(path)
    if suffix == ".docx":
        return _extract_docx(path)
    raise ExtractionError(f"unsupported_file_type:{suffix}")


def _decode_text(payload: bytes) -> str:
    for encoding in ("utf-8", "utf-8-sig", "latin-1"):
        try:
            return payload.decode(encoding).strip()
        except UnicodeDecodeError:
            continue
    raise ExtractionError("text_encoding_not_supported")


def _extract_csv(path: Path) -> ExtractedContent:
    text = _decode_text(path.read_bytes())
    rows = csv.reader(io.StringIO(text))
    normalized = [" | ".join(cell.strip() for cell in row) for row in rows]
    return ExtractedContent("\n".join(normalized).strip(), 1)


def _extract_json(path: Path) -> ExtractedContent:
    try:
        value = json.loads(_decode_text(path.read_bytes()))
    except json.JSONDecodeError as error:
        raise ExtractionError("invalid_json") from error
    return ExtractedContent(
        json.dumps(value, ensure_ascii=False, indent=2), 1
    )


def _extract_pdf(path: Path) -> ExtractedContent:
    try:
        reader = PdfReader(str(path))
        pages = [(page.extract_text() or "").strip() for page in reader.pages]
    except Exception as error:
        raise ExtractionError("invalid_pdf") from error
    text = "\n\n".join(page for page in pages if page).strip()
    if len(text) < 20:
        raise ExtractionError("ocr_required")
    return ExtractedContent(text, max(1, len(reader.pages)))


def _extract_docx(path: Path) -> ExtractedContent:
    try:
        document = WordDocument(str(path))
        paragraphs = [
            paragraph.text.strip()
            for paragraph in document.paragraphs
            if paragraph.text.strip()
        ]
        tables = []
        for table in document.tables:
            for row in table.rows:
                tables.append(
                    " | ".join(cell.text.strip() for cell in row.cells)
                )
    except Exception as error:
        raise ExtractionError("invalid_docx") from error
    text = "\n".join([*paragraphs, *tables]).strip()
    if len(text) < 20:
        raise ExtractionError("document_has_no_extractable_text")
    estimated_pages = max(1, round(len(text) / 2_800))
    return ExtractedContent(text, estimated_pages)
